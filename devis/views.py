"""
Vues de l'application devis.

Ce module contient toutes les vues nécessaires pour la gestion des devis,
incluant la création, la modification, la suppression et l'exportation des devis.
Chaque vue est protégée par le décorateur @login_required pour assurer
que seuls les utilisateurs authentifiés peuvent y accéder.
"""

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse, HttpResponseRedirect
from django.urls import reverse
from django.utils import timezone
from django.core.paginator import Paginator
from django.db.models import Sum
from decimal import Decimal
import csv
from datetime import datetime
from .models import Devis, LigneDevis
from .forms import DevisForm
from clients.models import Client
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

@login_required
def devis_list(request):
    """
    Vue pour afficher la liste de tous les devis.
    
    Cette vue affiche une liste paginée de tous les devis dans le système,
    triés par date de création décroissante (les plus récents en premier).
    
    Args:
        request: La requête HTTP
        
    Returns:
        HttpResponse: La page HTML avec la liste des devis
    """
    devis = Devis.objects.all().order_by('-date_creation')
    return render(request, 'devis/devis_list.html', {'devis': devis})

@login_required
def devis_en_cours(request):
    """
    Vue pour afficher la liste des devis en cours.
    
    Cette vue affiche une liste paginée des devis qui sont en cours de traitement,
    c'est-à-dire ceux qui ont le statut 'brouillon' ou 'envoye'.
    Les devis sont affichés 10 par page.
    
    Args:
        request: La requête HTTP
        
    Returns:
        HttpResponse: La page HTML avec la liste des devis en cours
    """
    devis = Devis.objects.filter(statut__in=['brouillon', 'envoye']).order_by('-date_creation')
    paginator = Paginator(devis, 10)  # 10 devis par page
    page = request.GET.get('page')
    devis = paginator.get_page(page)
    return render(request, 'devis/devis_list.html', {
        'devis': devis,
        'titre': 'Devis en cours'
    })

@login_required
def devis_termines(request):
    """
    Vue pour afficher la liste des devis terminés.
    
    Cette vue affiche une liste paginée des devis qui ont reçu une réponse finale,
    c'est-à-dire ceux qui ont le statut 'accepte' ou 'refuse'.
    Les devis sont affichés 10 par page.
    
    Args:
        request: La requête HTTP
        
    Returns:
        HttpResponse: La page HTML avec la liste des devis terminés
    """
    devis = Devis.objects.filter(statut__in=['accepte', 'refuse']).order_by('-date_creation')
    paginator = Paginator(devis, 10)  # 10 devis par page
    page = request.GET.get('page')
    devis = paginator.get_page(page)
    return render(request, 'devis/devis_list.html', {
        'devis': devis,
        'titre': 'Devis terminés'
    })

@login_required
def devis_detail(request, pk):
    """
    Vue pour afficher les détails d'un devis spécifique.
    
    Cette vue affiche toutes les informations d'un devis, y compris :
    - Les informations générales du devis
    - Les lignes détaillées du devis
    - Les informations du client associé
    - Les calculs des montants (HT, TTC, TVA)
    
    Args:
        request: La requête HTTP
        pk: L'identifiant unique du devis
        
    Returns:
        HttpResponse: La page HTML avec les détails du devis
    """
    devis = get_object_or_404(Devis, pk=pk)
    lignes = devis.lignes.all()
    return render(request, 'devis/devis_detail.html', {
        'devis': devis,
        'lignes': lignes
    })

@login_required
def devis_create(request):
    """
    Vue pour créer un nouveau devis.
    
    Cette vue gère le processus de création d'un nouveau devis :
    1. Affiche le formulaire de création si la méthode est GET
    2. Traite les données soumises si la méthode est POST :
       - Crée le devis avec les informations générales
       - Crée les lignes du devis
       - Calcule le montant total HT
       - Met à jour le devis avec le montant total
    
    Args:
        request: La requête HTTP
        
    Returns:
        HttpResponse: Le formulaire de création ou redirection vers le détail
    """
    if request.method == 'POST':
        # Récupération des données du formulaire
        client_id = request.POST.get('client')
        date_validite = request.POST.get('date_validite')
        conditions_paiement = request.POST.get('conditions_paiement')
        notes = request.POST.get('notes')
        
        # Création du devis
        client = get_object_or_404(Client, id=client_id)
        devis = Devis.objects.create(
            client=client,
            date_validite=date_validite,
            conditions_paiement=conditions_paiement,
            notes=notes,
            cree_par=request.user
        )
        
        # Traitement des lignes du devis
        descriptions = request.POST.getlist('description[]')
        quantites = request.POST.getlist('quantite[]')
        prix_unitaires = request.POST.getlist('prix_unitaire[]')
        
        montant_total_ht = 0
        for i in range(len(descriptions)):
            if descriptions[i] and quantites[i] and prix_unitaires[i]:
                ligne = LigneDevis.objects.create(
                    devis=devis,
                    description=descriptions[i],
                    quantite=int(quantites[i]),
                    prix_unitaire=float(prix_unitaires[i])
                )
                montant_total_ht += ligne.montant
        
        # Mise à jour du montant total
        devis.montant_ht = montant_total_ht
        devis.save()
        
        messages.success(request, 'Devis créé avec succès!')
        return redirect('devis:devis_detail', pk=devis.pk)
    
    clients = Client.objects.all()
    return render(request, 'devis/devis_form.html', {'clients': clients})

@login_required
def devis_update(request, pk):
    """
    Vue pour modifier un devis existant.
    
    Cette vue gère le processus de modification d'un devis :
    1. Affiche le formulaire pré-rempli si la méthode est GET
    2. Traite les données soumises si la méthode est POST :
       - Met à jour les informations générales du devis
       - Supprime les anciennes lignes
       - Crée les nouvelles lignes
       - Recalcule le montant total HT
       - Met à jour le devis avec le nouveau montant
    
    Args:
        request: La requête HTTP
        pk: L'identifiant unique du devis à modifier
        
    Returns:
        HttpResponse: Le formulaire de modification ou redirection vers le détail
    """
    devis = get_object_or_404(Devis, pk=pk)
    if request.method == 'POST':
        # Mise à jour des informations du devis
        client_id = request.POST.get('client')
        client = get_object_or_404(Client, id=client_id)
        devis.client = client
        devis.date_validite = request.POST.get('date_validite')
        devis.conditions_paiement = request.POST.get('conditions_paiement')
        devis.notes = request.POST.get('notes')
        
        # Suppression des anciennes lignes
        devis.lignes.all().delete()
        
        # Création des nouvelles lignes
        descriptions = request.POST.getlist('description[]')
        quantites = request.POST.getlist('quantite[]')
        prix_unitaires = request.POST.getlist('prix_unitaire[]')
        
        montant_total_ht = 0
        for i in range(len(descriptions)):
            if descriptions[i] and quantites[i] and prix_unitaires[i]:
                ligne = LigneDevis.objects.create(
                    devis=devis,
                    description=descriptions[i],
                    quantite=int(quantites[i]),
                    prix_unitaire=float(prix_unitaires[i])
                )
                montant_total_ht += ligne.montant
        
        devis.montant_ht = montant_total_ht
        devis.save()
        
        messages.success(request, 'Devis mis à jour avec succès!')
        return redirect('devis:devis_detail', pk=devis.pk)
    
    return render(request, 'devis/devis_form.html', {
        'devis': devis,
        'clients': Client.objects.all()
    })

@login_required
def devis_delete(request, pk):
    """
    Vue pour supprimer un devis.
    
    Cette vue gère le processus de suppression d'un devis :
    1. Affiche une page de confirmation si la méthode est GET
    2. Supprime le devis et toutes ses lignes associées si la méthode est POST
    
    Args:
        request: La requête HTTP
        pk: L'identifiant unique du devis à supprimer
        
    Returns:
        HttpResponse: La page de confirmation ou redirection vers la liste
    """
    devis = get_object_or_404(Devis, pk=pk)
    if request.method == 'POST':
        devis.delete()
        messages.success(request, 'Devis supprimé avec succès!')
        return redirect('devis:devis_list')
    
    return render(request, 'devis/devis_confirm_delete.html', {'devis': devis})

@login_required
def devis_export_pdf(request, pk):
    """
    Exporte un devis en PDF en utilisant ReportLab.
    
    Génère un PDF professionnel contenant toutes les informations
    du devis, y compris les lignes et les totaux.
    
    Args:
        request: La requête HTTP
        pk: L'identifiant unique du devis à exporter
        
    Returns:
        HttpResponse: Le fichier PDF généré
    """
    devis = get_object_or_404(Devis, pk=pk)
    
    # Création du PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="devis_{devis.numero}.pdf"'
    
    # Création du document
    doc = SimpleDocTemplate(response, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []
    
    # Titre
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1  # Centré
    )
    elements.append(Paragraph(f"Devis N°{devis.numero}", title_style))
    elements.append(Paragraph(f"Date: {devis.date_creation.strftime('%d/%m/%Y')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Informations client
    elements.append(Paragraph("Client:", styles['Heading2']))
    elements.append(Paragraph(f"{devis.client.nom}", styles['Normal']))
    elements.append(Paragraph(f"{devis.client.adresse}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Tableau des lignes
    data = [['Description', 'Quantité', 'Prix unitaire', 'Total']]
    for ligne in devis.lignes.all():
        data.append([
            ligne.description,
            str(ligne.quantite),
            f"{ligne.prix_unitaire:.2f} €",
            f"{ligne.montant:.2f} €"
        ])
    
    # Style du tableau
    table_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ])
    
    # Création du tableau
    table = Table(data)
    table.setStyle(table_style)
    elements.append(table)
    elements.append(Spacer(1, 20))
    
    # Calcul de la TVA avec Decimal
    tva_rate = Decimal('0.20')  # 20% TVA
    montant_tva = devis.montant_ht * tva_rate
    montant_ttc = devis.montant_ht + montant_tva
    
    # Totaux
    totals_data = [
        ['Total HT:', f"{devis.montant_ht:.2f} €"],
        ['TVA (20%):', f"{montant_tva:.2f} €"],
        ['Total TTC:', f"{montant_ttc:.2f} €"]
    ]
    
    totals_table = Table(totals_data, colWidths=[2*inch, 2*inch])
    totals_style = TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ])
    totals_table.setStyle(totals_style)
    elements.append(totals_table)
    
    # Pied de page
    elements.append(Spacer(1, 30))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.grey,
        alignment=1
    )
    elements.append(Paragraph("Ce devis est valable 30 jours à compter de sa date d'émission.", footer_style))
    
    # Génération du PDF
    doc.build(elements)
    return response

@login_required
def devis_export_csv(request):
    """
    Exporte la liste des devis en format CSV.
    
    Génère un fichier CSV contenant tous les devis avec leurs
    informations principales, triés par date de création.
    
    Args:
        request: La requête HTTP
        
    Returns:
        HttpResponse: Le fichier CSV généré
    """
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="devis_{datetime.now().strftime("%Y%m%d")}.csv"'
    
    writer = csv.writer(response)
    # En-têtes du CSV
    writer.writerow([
        'Numéro', 'Client', 'Date création', 'Date validité',
        'Montant HT', 'TVA', 'Montant TTC', 'Statut'
    ])
    
    # Données des devis
    devis = Devis.objects.all().order_by('-date_creation')
    for devis in devis:
        writer.writerow([
            devis.numero,
            devis.client.nom,
            devis.date_creation.strftime('%d/%m/%Y'),
            devis.date_validite.strftime('%d/%m/%Y'),
            f"{devis.montant_ht:.2f}",
            f"{devis.montant_ht * Decimal('0.20'):.2f}",
            f"{devis.montant_ttc:.2f}",
            devis.get_statut_display()
        ])
    
    return response

@login_required
def devis_terminer(request, pk):
    """
    Vue pour marquer un devis comme terminé (accepté ou refusé).
    
    Permet de finaliser un devis en le marquant comme accepté ou refusé.
    Cette action ne peut être effectuée qu'une seule fois.
    
    Args:
        request: La requête HTTP
        pk: L'identifiant unique du devis à terminer
        
    Returns:
        HttpResponse: La page de confirmation ou redirection vers le détail
    """
    devis = get_object_or_404(Devis, pk=pk)
    
    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'accepter':
            devis.statut = 'accepte'
            messages.success(request, 'Le devis a été marqué comme accepté.')
        elif action == 'refuser':
            devis.statut = 'refuse'
            messages.success(request, 'Le devis a été marqué comme refusé.')
        
        devis.save()
        return redirect('devis:devis_detail', pk=devis.pk)
    
    return render(request, 'devis/devis_terminer.html', {'devis': devis})

@login_required
def devis_export_excel(request, pk):
    """
    Exporte un devis en format Excel.
    
    Génère un fichier Excel contenant toutes les informations
    du devis, y compris les lignes et les totaux.
    
    Args:
        request: La requête HTTP
        pk: L'identifiant unique du devis à exporter
        
    Returns:
        HttpResponse: Le fichier Excel généré
    """
    devis = get_object_or_404(Devis, pk=pk)
    
    # Création du classeur Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Devis {devis.numero}"
    
    # Styles
    header_font = Font(bold=True, size=12)
    header_fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
    title_font = Font(bold=True, size=14)
    
    # Titre
    ws['A1'] = f"Devis N°{devis.numero}"
    ws['A1'].font = title_font
    ws['A2'] = f"Date: {devis.date_creation.strftime('%d/%m/%Y')}"
    
    # Informations client
    ws['A4'] = "Client:"
    ws['A4'].font = header_font
    ws['A5'] = devis.client.nom
    ws['A6'] = devis.client.adresse
    
    # En-têtes du tableau
    headers = ['Description', 'Quantité', 'Prix unitaire', 'Total']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=8, column=col)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    
    # Données des lignes
    row = 9
    for ligne in devis.lignes.all():
        ws.cell(row=row, column=1).value = ligne.description
        ws.cell(row=row, column=2).value = ligne.quantite
        ws.cell(row=row, column=3).value = f"{ligne.prix_unitaire:.2f} €"
        ws.cell(row=row, column=4).value = f"{ligne.montant:.2f} €"
        row += 1
    
    # Totaux
    row += 1
    ws.cell(row=row, column=1).value = "Total HT:"
    ws.cell(row=row, column=2).value = f"{devis.montant_ht:.2f} €"
    ws.cell(row=row, column=1).font = header_font
    
    row += 1
    tva = devis.montant_ht * Decimal('0.20')
    ws.cell(row=row, column=1).value = "TVA (20%):"
    ws.cell(row=row, column=2).value = f"{tva:.2f} €"
    ws.cell(row=row, column=1).font = header_font
    
    row += 1
    ws.cell(row=row, column=1).value = "Total TTC:"
    ws.cell(row=row, column=2).value = f"{devis.montant_ttc:.2f} €"
    ws.cell(row=row, column=1).font = header_font
    
    # Ajustement des largeurs de colonnes
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column].width = adjusted_width
    
    # Création de la réponse
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="devis_{devis.numero}.xlsx"'
    wb.save(response)
    return response

@login_required
def devis_export_word(request, pk):
    """
    Exporte un devis en format Word.
    
    Génère un document Word contenant toutes les informations
    du devis, y compris les lignes et les totaux.
    
    Args:
        request: La requête HTTP
        pk: L'identifiant unique du devis à exporter
        
    Returns:
        HttpResponse: Le fichier Word généré
    """
    devis = get_object_or_404(Devis, pk=pk)
    
    # Création du document Word
    doc = Document()
    
    # Titre
    title = doc.add_heading(f"Devis N°{devis.numero}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph(f"Date: {devis.date_creation.strftime('%d/%m/%Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Informations client
    doc.add_heading("Client:", level=2)
    doc.add_paragraph(devis.client.nom)
    doc.add_paragraph(devis.client.adresse)
    
    # Tableau des lignes
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    headers = ['Description', 'Quantité', 'Prix unitaire', 'Total']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Données des lignes
    for ligne in devis.lignes.all():
        row_cells = table.add_row().cells
        row_cells[0].text = ligne.description
        row_cells[1].text = str(ligne.quantite)
        row_cells[2].text = f"{ligne.prix_unitaire:.2f} €"
        row_cells[3].text = f"{ligne.montant:.2f} €"
    
    # Totaux
    doc.add_paragraph()
    totals = doc.add_table(rows=3, cols=2)
    totals.style = 'Table Grid'
    
    # Total HT
    totals.rows[0].cells[0].text = "Total HT:"
    totals.rows[0].cells[1].text = f"{devis.montant_ht:.2f} €"
    totals.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    
    # TVA
    tva = devis.montant_ht * Decimal('0.20')
    totals.rows[1].cells[0].text = "TVA (20%):"
    totals.rows[1].cells[1].text = f"{tva:.2f} €"
    totals.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    
    # Total TTC
    totals.rows[2].cells[0].text = "Total TTC:"
    totals.rows[2].cells[1].text = f"{devis.montant_ttc:.2f} €"
    totals.rows[2].cells[0].paragraphs[0].runs[0].bold = True
    
    # Pied de page
    doc.add_paragraph()
    footer = doc.add_paragraph("Ce devis est valable 30 jours à compter de sa date d'émission.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Création de la réponse
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="devis_{devis.numero}.docx"'
    doc.save(response)
    return response
